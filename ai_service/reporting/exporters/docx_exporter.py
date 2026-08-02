"""Word (.docx) exporter — builds a Word document from report data using python-docx."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
from datetime import datetime, timezone


def export_docx(title: str, report_type_label: str, session_id: str,
                session_start, session_end, summary: dict,
                ai_narrative: str | None, flows: list,
                alerts: list, output_path: str):
    """Generate a .docx report and write it to output_path."""

    doc = Document()

    # ── Title page ────────────────────────────────────────────────────────────
    heading = doc.add_heading('', 0)
    run = heading.add_run('PeliCap')
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # blue-600
    run.font.size = Pt(14)

    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_table = doc.add_table(rows=4, cols=2)
    meta_data = [
        ("Session ID", str(session_id)),
        ("Period", f"{session_start} — {session_end}"),
        ("Report Type", report_type_label),
        ("Generated", datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v

    doc.add_page_break()

    # ── Executive summary ─────────────────────────────────────────────────────
    if ai_narrative:
        doc.add_heading('Executive Summary', level=2)
        for para in ai_narrative.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_page_break()

    # ── Session statistics ────────────────────────────────────────────────────
    doc.add_heading('Session Statistics', level=2)
    stats_table = doc.add_table(rows=6, cols=2)
    stats_table.style = 'Table Grid'
    stats = [
        ("Total Flows", f"{summary.get('total_flows', 0):,}"),
        ("Data Transferred", _fmt_bytes(summary.get('total_bytes', 0))),
        ("Average RTT", f"{summary.get('avg_rtt_us', 0) / 1000:.1f} ms"),
        ("Retransmits", f"{summary.get('total_retransmits', 0):,}"),
        ("Unique Sources", f"{summary.get('unique_src_ips', 0):,}"),
        ("Unique Destinations", f"{summary.get('unique_dst_ips', 0):,}"),
    ]
    for i, (k, v) in enumerate(stats):
        stats_table.rows[i].cells[0].text = k
        stats_table.rows[i].cells[1].text = v

    # Protocol breakdown
    proto_bd = summary.get("protocol_breakdown", [])
    if proto_bd:
        doc.add_heading('Protocol Breakdown', level=3)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = 'Protocol'; hdr[1].text = 'Flows'; hdr[2].text = 'Data'
        for p in proto_bd:
            row = tbl.add_row().cells
            row[0].text = p['protocol']
            row[1].text = f"{p['count']:,}"
            row[2].text = _fmt_bytes(p['bytes'])

    # ── Alerts ────────────────────────────────────────────────────────────────
    if alerts:
        doc.add_page_break()
        doc.add_heading(f'Security Alerts ({len(alerts)})', level=2)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        hdr[0].text = 'Severity'; hdr[1].text = 'Title'
        hdr[2].text = 'Description'; hdr[3].text = 'Source IP'
        for a in alerts[:50]:
            row = tbl.add_row().cells
            row[0].text = str(a.get('severity', ''))
            row[1].text = str(a.get('title', ''))
            row[2].text = str(a.get('description', ''))[:100]
            row[3].text = str(a.get('src_ip') or '—')

    # ── Flows appendix ────────────────────────────────────────────────────────
    if flows:
        doc.add_page_break()
        doc.add_heading(f'Flow Appendix (top {min(100, len(flows))})', level=2)
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        headers = ['Source', 'Destination', 'Protocol', 'Bytes', 'RTT', 'Retransmits']
        for i, h in enumerate(headers):
            hdr[i].text = h
        proto_map = {6: 'TCP', 17: 'UDP', 1: 'ICMP'}
        for f in flows[:100]:
            row = tbl.add_row().cells
            row[0].text = f"{f.get('src_ip')}:{f.get('src_port')}"
            row[1].text = f"{f.get('dst_ip')}:{f.get('dst_port')}"
            row[2].text = proto_map.get(f.get('protocol', 0), '?')
            b = (f.get('fwd_bytes') or 0) + (f.get('rev_bytes') or 0)
            row[3].text = _fmt_bytes(b)
            row[4].text = f"{(f.get('avg_rtt_us') or 0) / 1000:.1f}ms"
            row[5].text = str(f.get('retransmit_count') or 0)

    doc.save(output_path)


def _fmt_bytes(b: int | float) -> str:
    b = b or 0
    if b >= 1e9: return f"{b/1e9:.2f} GB"
    if b >= 1e6: return f"{b/1e6:.2f} MB"
    if b >= 1e3: return f"{b/1e3:.1f} KB"
    return f"{b:.0f} B"
